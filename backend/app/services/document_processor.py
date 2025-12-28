"""Document processing service for PDFs, DOCX, and text files."""

import io
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# PDF processing
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

# DOCX processing
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocumentProcessor:
    """Process various document formats for RAG ingestion."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = [".pdf", ".txt", ".md"]
        if HAS_DOCX:
            self.supported_formats.append(".docx")

    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        metadata: Optional[Dict] = None
    ) -> Dict:
        """
        Process a file and extract text content.

        Args:
            file_content: Raw file bytes
            filename: Original filename
            metadata: Additional metadata

        Returns:
            Dict with title, content, chunks, and metadata
        """
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return await self._process_pdf(file_content, filename, metadata)
        elif ext == ".docx":
            return await self._process_docx(file_content, filename, metadata)
        elif ext in [".txt", ".md"]:
            return await self._process_text(file_content, filename, metadata)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    async def _process_pdf(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict]
    ) -> Dict:
        """Extract text from PDF using PyMuPDF."""
        if not HAS_PYMUPDF:
            raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")

        try:
            # Open PDF from bytes
            doc = fitz.open(stream=content, filetype="pdf")

            # Extract metadata
            pdf_metadata = doc.metadata or {}
            title = pdf_metadata.get("title") or Path(filename).stem
            author = pdf_metadata.get("author", "")
            creation_date = pdf_metadata.get("creationDate", "")

            # Extract text from all pages
            full_text = ""
            page_texts = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text("text")
                page_texts.append({
                    "page": page_num + 1,
                    "text": page_text
                })
                full_text += page_text + "\n\n"

            doc.close()

            # Clean up text
            full_text = self._clean_text(full_text)

            # Create chunks
            chunks = self._create_chunks(full_text)

            return {
                "success": True,
                "title": title,
                "content": full_text,
                "chunks": chunks,
                "page_count": len(page_texts),
                "pages": page_texts,
                "format": "pdf",
                "metadata": {
                    **(metadata or {}),
                    "author": author,
                    "creation_date": creation_date,
                    "filename": filename,
                }
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    async def _process_docx(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict]
    ) -> Dict:
        """Extract text from DOCX file."""
        if not HAS_DOCX:
            raise ImportError("python-docx not installed. Run: pip install python-docx")

        try:
            # Open DOCX from bytes
            doc = DocxDocument(io.BytesIO(content))

            # Extract title from first heading or filename
            title = Path(filename).stem
            for para in doc.paragraphs[:5]:
                if para.style.name.startswith("Heading"):
                    title = para.text
                    break

            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n\n".join(paragraphs)
            full_text = self._clean_text(full_text)

            # Create chunks
            chunks = self._create_chunks(full_text)

            return {
                "success": True,
                "title": title,
                "content": full_text,
                "chunks": chunks,
                "paragraph_count": len(paragraphs),
                "format": "docx",
                "metadata": {
                    **(metadata or {}),
                    "filename": filename,
                }
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    async def _process_text(
        self,
        content: bytes,
        filename: str,
        metadata: Optional[Dict]
    ) -> Dict:
        """Process plain text or markdown files."""
        try:
            # Decode text
            text = content.decode("utf-8")
            text = self._clean_text(text)

            # Extract title from first line or filename
            lines = text.split("\n")
            title = Path(filename).stem
            if lines and lines[0].startswith("#"):
                title = lines[0].lstrip("#").strip()
            elif lines:
                title = lines[0][:100]

            # Create chunks
            chunks = self._create_chunks(text)

            return {
                "success": True,
                "title": title,
                "content": text,
                "chunks": chunks,
                "format": "text",
                "metadata": {
                    **(metadata or {}),
                    "filename": filename,
                }
            }

        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }

    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        import re

        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

        return text.strip()

    def _create_chunks(self, text: str) -> List[Dict]:
        """Split text into overlapping chunks."""
        if len(text) <= self.chunk_size:
            return [{
                "content": text,
                "chunk_index": 0,
                "char_start": 0,
                "char_end": len(text)
            }]

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence end in last 20% of chunk
                search_start = end - int(self.chunk_size * 0.2)
                best_break = end

                for i in range(end - 1, search_start - 1, -1):
                    if text[i] in '.!?\n':
                        best_break = i + 1
                        break

                end = best_break

            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append({
                    "content": chunk_text,
                    "chunk_index": chunk_index,
                    "char_start": start,
                    "char_end": end
                })
                chunk_index += 1

            # Move start with overlap
            start = end - self.chunk_overlap
            if start >= len(text) or start < 0:
                break

        return chunks

    async def extract_images_from_pdf(
        self,
        content: bytes,
        max_images: int = 10
    ) -> List[Dict]:
        """Extract images from PDF for vision processing."""
        if not HAS_PYMUPDF:
            return []

        try:
            doc = fitz.open(stream=content, filetype="pdf")
            images = []

            for page_num, page in enumerate(doc):
                image_list = page.get_images(full=True)

                for img_index, img in enumerate(image_list):
                    if len(images) >= max_images:
                        break

                    xref = img[0]
                    base_image = doc.extract_image(xref)

                    images.append({
                        "page": page_num + 1,
                        "index": img_index,
                        "width": base_image.get("width"),
                        "height": base_image.get("height"),
                        "format": base_image.get("ext"),
                        "data": base_image.get("image"),  # Raw bytes
                    })

            doc.close()
            return images

        except Exception as e:
            print(f"Image extraction error: {e}")
            return []


# Singleton instance
document_processor = DocumentProcessor()
